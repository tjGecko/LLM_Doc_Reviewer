"""DOCX document loader with stable paragraph ID generation."""

import hashlib
import re
from pathlib import Path
from typing import List, Dict, Optional

try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

from ..config import DocumentChunk


def load_paragraphs(path: str) -> List[DocumentChunk]:
    """
    Load paragraphs from a DOCX file with stable paragraph IDs.
    
    Args:
        path: Path to the DOCX file
        
    Returns:
        List of DocumentChunk objects with stable IDs
        
    Raises:
        ImportError: If python-docx is not installed
        Exception: If DOCX cannot be read
    """
    if not HAS_DOCX:
        raise ImportError(
            "python-docx is required for DOCX support. Install with: pip install python-docx"
        )
    
    file_path = Path(path)
    chunks = []
    chunk_index = 0
    
    try:
        doc = docx.Document(file_path)
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            
            # Skip empty paragraphs
            if not text:
                continue
            
            # Skip very short paragraphs (likely formatting artifacts)
            if len(text) < 10:
                continue
            
            # Clean the text
            text = clean_docx_text(text)
            
            # Generate stable paragraph ID
            content_hash = hashlib.md5(text.encode('utf-8')).hexdigest()[:8]
            paragraph_id = f"{file_path.stem}_{content_hash}_{chunk_index:03d}"
            
            # Create document chunk
            chunk = DocumentChunk(
                paragraph_id=paragraph_id,
                text=text,
                hash=content_hash,
                page_number=None,  # DOCX doesn't have reliable page numbers
                chunk_index=chunk_index
            )
            
            chunks.append(chunk)
            chunk_index += 1
    
    except Exception as e:
        raise Exception(f"Failed to read DOCX {path}: {e}")
    
    return chunks


def clean_docx_text(text: str) -> str:
    """
    Clean common DOCX extraction artifacts.
    
    Args:
        text: Raw text extracted from DOCX
        
    Returns:
        Cleaned text
    """
    # Normalize whitespace
    text = re.sub(r'\s+', ' ', text)
    
    # Remove common artifacts
    # Remove tab characters
    text = text.replace('\t', ' ')
    
    # Clean up multiple spaces
    text = re.sub(r' +', ' ', text)
    
    return text.strip()


def is_supported_format(path: str) -> bool:
    """Check if the file format is supported by this loader."""
    supported_extensions = {'.docx', '.doc'}
    return Path(path).suffix.lower() in supported_extensions


def get_document_info(path: str) -> Dict[str, any]:
    """Get basic information about the DOCX document."""
    if not HAS_DOCX:
        return {
            'format': 'docx',
            'error': 'python-docx not installed'
        }
    
    file_path = Path(path)
    
    try:
        doc = docx.Document(file_path)
        
        # Extract core properties
        props = doc.core_properties
        
        # Count paragraphs and estimate content
        paragraph_count = 0
        word_count = 0
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text and len(text) >= 10:  # Only count substantial paragraphs
                paragraph_count += 1
                word_count += len(text.split())
        
        # Extract document statistics
        return {
            'format': 'docx',
            'file_size': file_path.stat().st_size,
            'paragraph_count': paragraph_count,
            'word_count': word_count,
            'title': props.title or 'Unknown',
            'author': props.author or 'Unknown',
            'subject': props.subject or 'Unknown',
            'created': str(props.created) if props.created else 'Unknown',
            'modified': str(props.modified) if props.modified else 'Unknown',
            'last_modified_by': props.last_modified_by or 'Unknown'
        }
    
    except Exception as e:
        return {
            'format': 'docx',
            'file_size': file_path.stat().st_size,
            'error': str(e)
        }